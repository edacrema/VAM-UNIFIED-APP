from __future__ import annotations

import base64
import io
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from .report_blocks import ReportBlock


def _safe_filename(filename: str) -> str:
    name = filename.strip() or "export.docx"
    name = name.replace("\\", "_").replace("/", "_")
    name = re.sub(r"[^A-Za-z0-9._\- ]+", "_", name)
    if not name.lower().endswith(".docx"):
        name = name + ".docx"
    return name


def _add_text_lines(doc: Document, text: str) -> None:
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue

        doc.add_paragraph(line)


def _get_risk_color_rgb(score: float) -> tuple[int, int, int]:
    if score < 4.0:
        return (214, 39, 40)
    if score < 5.5:
        return (255, 127, 14)
    if score < 7.0:
        return (255, 187, 120)
    return (44, 160, 44)


def _set_cell_background(cell: Any, rgb_tuple: tuple[int, int, int]) -> None:
    r, g, b = rgb_tuple
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{r:02x}{g:02x}{b:02x}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_overview_table_to_document(doc: Document, *, meta: Dict[str, Any]) -> None:
    dims = meta.get("dimensions") or []
    rows = meta.get("rows") or []
    if not isinstance(dims, list) or not isinstance(rows, list) or not dims or not rows:
        return

    sorted_rows = []
    for r in rows:
        if not isinstance(r, dict):
            continue
        sorted_rows.append(r)
    sorted_rows = sorted(sorted_rows, key=lambda x: float(x.get("overall_mfi", 0) or 0))

    n_rows = len(sorted_rows) + 1
    n_cols = len(dims) + 3

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Market", "Region"] + [str(d) for d in dims] + ["MFI"]
    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = str(header)
        _set_cell_background(cell, (0, 114, 188))
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, market in enumerate(sorted_rows, start=1):
        row = table.rows[row_idx]

        market_name = str(market.get("market_name", "") or "").strip()
        row.cells[0].text = market_name

        region = str(market.get("region", "") or "").strip()
        row.cells[1].text = region

        dim_scores = market.get("dimension_scores")
        if not isinstance(dim_scores, dict):
            dim_scores = {}

        for dim_idx, dim in enumerate(dims):
            cell = row.cells[dim_idx + 2]
            try:
                score = float(dim_scores.get(dim, 0) or 0)
            except Exception:
                score = 0.0
            cell.text = f"{score:.1f}"
            _set_cell_background(cell, _get_risk_color_rgb(score))

            text_color = RGBColor(255, 255, 255) if score < 4.0 or score >= 7.0 else RGBColor(0, 0, 0)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = text_color

        mfi_cell = row.cells[-1]
        try:
            mfi_score = float(market.get("overall_mfi", 0) or 0)
        except Exception:
            mfi_score = 0.0
        mfi_cell.text = f"{mfi_score:.1f}"
        _set_cell_background(mfi_cell, _get_risk_color_rgb(mfi_score))
        mfi_text_color = RGBColor(255, 255, 255) if mfi_score < 4.0 or mfi_score >= 7.0 else RGBColor(0, 0, 0)
        for paragraph in mfi_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = mfi_text_color
                run.bold = True

        for cell_idx in [0, 1]:
            cell = row.cells[cell_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()


def _add_definition_box(doc: Document, text: str) -> None:
    cleaned = (text or "").strip()
    if not cleaned:
        return

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"

    cell = table.rows[0].cells[0]
    header_para = cell.paragraphs[0]

    header_run = header_para.add_run("Definition: ")
    header_run.bold = True
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0, 114, 188)

    def_run = header_para.add_run(cleaned)
    def_run.italic = True
    def_run.font.size = Pt(9)

    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6F3FF"/>')
    cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()


def build_docx_bytes_from_report_blocks(
    report_blocks: List[ReportBlock],
    *,
    visualizations: Optional[Dict[str, str]] = None,
    include_sources: bool = True,
    include_visualizations: bool = True,
) -> bytes:
    doc = Document()
    visualizations = visualizations or {}

    for block in report_blocks:
        if block.type == "heading":
            level = int(block.level or 1)
            level = min(max(level, 1), 9)
            doc.add_heading(block.text or "", level=level)
            continue

        if block.type == "paragraph":
            _add_text_lines(doc, block.text or "")
            continue

        if block.type == "figure":
            if not include_visualizations:
                continue

            fig_id = (block.figure_id or "").strip()
            fig_b64 = visualizations.get(fig_id)
            if not fig_id or not fig_b64:
                continue

            try:
                img_bytes = base64.b64decode(fig_b64)
            except Exception:
                continue

            width = float(block.width) if block.width is not None else 6.0
            width = max(1.0, min(width, 7.0))

            buf = io.BytesIO(img_bytes)
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(buf, width=Inches(width))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if block.caption:
                cap = doc.add_paragraph(block.caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if block.type == "references":
            if not include_sources:
                continue

            refs = block.references or []
            if not refs:
                continue

            doc.add_heading("References", level=2)
            for ref in refs:
                if not isinstance(ref, dict):
                    continue

                doc_id = str(ref.get("doc_id", "")).strip()
                source = str(ref.get("source", "")).strip()
                date = str(ref.get("date", "")).strip()
                title = str(ref.get("title", "")).strip()
                url = str(ref.get("url", "")).strip()

                parts: List[str] = []
                if doc_id:
                    parts.append(f"[{doc_id}]")
                if source:
                    parts.append(source)
                if date:
                    parts.append(f"({date})")
                if title:
                    parts.append(title)

                doc.add_paragraph(" ".join(parts).strip(), style="List Number")
                if url:
                    doc.add_paragraph(url)
            continue

        if block.type == "table":
            meta = block.meta or {}
            if isinstance(meta, dict) and meta.get("table_kind") == "mfi_overview":
                _add_overview_table_to_document(doc, meta=meta)
            continue

        if block.type == "definition_box":
            _add_definition_box(doc, block.text or "")
            continue

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def build_content_disposition(filename: str) -> str:
    safe = _safe_filename(filename)
    return f'attachment; filename="{safe}"'
